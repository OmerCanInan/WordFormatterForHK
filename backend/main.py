"""
WordFormatter – FastAPI Backend

Provides REST endpoints for uploading two Word documents (template + content)
and returning a merged result. Designed for extensibility (AI pipeline hooks).
"""

from __future__ import annotations

import io
import logging
from pathlib import Path

from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from docx import Document
from dotenv import load_dotenv

from merger import merge_documents
from ai_merger import ai_merge_documents

# Çevre değişkenlerini (API Key vb.) yükle
load_dotenv()

# ---------------------------------------------------------------------------
# App Setup
# ---------------------------------------------------------------------------

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("wordformatter")

app = FastAPI(
    title="WordFormatter API",
    description="Merge Word document formatting and content",
    version="1.0.0",
)

# CORS – allow Vite dev server
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ---------------------------------------------------------------------------
# Health Check
# ---------------------------------------------------------------------------

@app.get("/api/health")
async def health_check():
    return {"status": "ok", "version": "1.0.0"}


# ---------------------------------------------------------------------------
# Merge Endpoint
# ---------------------------------------------------------------------------

ALLOWED_CONTENT_TYPES = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/octet-stream",  # some browsers send this
}

ALLOWED_EXTENSIONS = {".docx"}


def _validate_upload(file: UploadFile, label: str) -> None:
    """Validate that the uploaded file is a .docx."""
    if file.filename:
        ext = Path(file.filename).suffix.lower()
        if ext not in ALLOWED_EXTENSIONS:
            raise HTTPException(
                status_code=400,
                detail=f"{label}: Sadece .docx dosyaları kabul edilmektedir. "
                       f"Yüklenen dosya uzantısı: {ext}",
            )


@app.post("/api/merge")
async def merge_files(
    template: UploadFile = File(..., description="Biçimlendirme Şablonu Word Dosyası"),
    content: UploadFile = File(..., description="Yazı İçeriği Word Dosyası"),
):
    """
    Accepts two .docx files:
      - template: provides the formatting/styles
      - content:  provides the raw text

    Returns a merged .docx file as a download.
    """
    # Validate uploads
    _validate_upload(template, "Şablon dosyası")
    _validate_upload(content, "İçerik dosyası")

    try:
        template_bytes = await template.read()
        content_bytes = await content.read()

        template_doc = Document(io.BytesIO(template_bytes))
        content_doc = Document(io.BytesIO(content_bytes))
    except Exception as exc:
        logger.error("Dosya okuma hatası: %s", exc)
        raise HTTPException(
            status_code=400,
            detail=f"Dosyalar okunurken hata oluştu: {str(exc)}",
        )

    try:
        result_doc = merge_documents(template_doc, content_doc)
    except Exception as exc:
        logger.error("Birleştirme hatası: %s", exc)
        raise HTTPException(
            status_code=500,
            detail=f"Dosyalar birleştirilirken hata oluştu: {str(exc)}",
        )

    # Serialize result to bytes
    output_buffer = io.BytesIO()
    result_doc.save(output_buffer)
    output_buffer.seek(0)

    return StreamingResponse(
        output_buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": 'attachment; filename="merged_output.docx"',
        },
    )


# ---------------------------------------------------------------------------
# Future: AI Pipeline Hooks (placeholder)
# ---------------------------------------------------------------------------

@app.post("/api/ai-merge")
async def ai_merge_files(
    template: UploadFile = File(..., description="Biçimlendirme Şablonu Word Dosyası"),
    content: UploadFile = File(..., description="Yazı İçeriği Word Dosyası"),
):
    """
    Accepts two .docx files and uses Gemini AI to intelligently
    merge content into the template's structure.
    """
    # Validate uploads
    _validate_upload(template, "Şablon dosyası")
    _validate_upload(content, "İçerik dosyası")

    try:
        template_bytes = await template.read()
        content_bytes = await content.read()

        template_doc = Document(io.BytesIO(template_bytes))
        content_doc = Document(io.BytesIO(content_bytes))
    except Exception as exc:
        logger.error("Dosya okuma hatası: %s", exc)
        raise HTTPException(
            status_code=400,
            detail=f"Dosyalar okunurken hata oluştu: {str(exc)}",
        )

    try:
        result_doc = ai_merge_documents(template_doc, content_doc)
    except Exception as exc:
        logger.error("AI Birleştirme hatası: %s", exc)
        raise HTTPException(
            status_code=500,
            detail=f"Dosyalar AI ile birleştirilirken hata oluştu: {str(exc)}",
        )

    # Serialize result to bytes
    output_buffer = io.BytesIO()
    result_doc.save(output_buffer)
    output_buffer.seek(0)

    return StreamingResponse(
        output_buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": 'attachment; filename="ai_merged_output.docx"',
        },
    )
