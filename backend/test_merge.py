"""
End-to-end test: Creates sample template + content docs, merges them,
and verifies the output.
"""
import io
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from merger import merge_documents

def create_template():
    """Template with rich formatting but placeholder text."""
    doc = Document()

    # Paragraph 1: Big bold blue title
    p1 = doc.add_paragraph()
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p1.add_run("SABLON BASLIK METNI")
    run1.font.name = "Arial"
    run1.font.size = Pt(24)
    run1.font.bold = True
    run1.font.color.rgb = RGBColor(0, 0, 128)  # Navy

    # Paragraph 2: Italic gray body
    p2 = doc.add_paragraph()
    run2 = p2.add_run("Sablon govde metni burada yer alir.")
    run2.font.name = "Times New Roman"
    run2.font.size = Pt(11)
    run2.font.italic = True
    run2.font.color.rgb = RGBColor(128, 128, 128)  # Gray

    # Paragraph 3: Small red footer
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run3 = p3.add_run("Sablon alt bilgi")
    run3.font.name = "Calibri"
    run3.font.size = Pt(9)
    run3.font.color.rgb = RGBColor(200, 0, 0)  # Red

    return doc

def create_content():
    """Content document with plain text."""
    doc = Document()
    doc.add_paragraph("Proje Durum Raporu - Mayis 2026")
    doc.add_paragraph("Bu ay icerisinde tum hedefler basariyla tamamlandi ve sistem uretim ortamina deploy edildi.")
    doc.add_paragraph("Hazirlayan: Omer - 21.05.2026")
    return doc

def main():
    template_doc = create_template()
    content_doc = create_content()

    print("=== SABLON PARAGRAFLARI ===")
    for i, p in enumerate(template_doc.paragraphs):
        runs_info = []
        for r in p.runs:
            runs_info.append(f"  font={r.font.name}, size={r.font.size}, bold={r.font.bold}, italic={r.font.italic}, color={r.font.color.rgb}")
        print(f"  P{i}: '{p.text}'")
        for info in runs_info:
            print(f"    {info}")

    print("\n=== ICERIK PARAGRAFLARI ===")
    for i, p in enumerate(content_doc.paragraphs):
        print(f"  P{i}: '{p.text}'")

    # Merge
    result = merge_documents(template_doc, content_doc)

    print("\n=== BIRLESTIRILMIS SONUC ===")
    for i, p in enumerate(result.paragraphs):
        runs_info = []
        for r in p.runs:
            runs_info.append(f"font={r.font.name}, size={r.font.size}, bold={r.font.bold}, italic={r.font.italic}, color={r.font.color.rgb if r.font.color and r.font.color.rgb else 'N/A'}")
        print(f"  P{i}: '{p.text}'")
        for info in runs_info:
            print(f"    {info}")

    # Save to verify
    result.save("test_output.docx")
    print("\n[OK] test_output.docx basariyla kaydedildi!")

    # Verify
    assert result.paragraphs[0].text == "Proje Durum Raporu - Mayis 2026", f"P0 text mismatch: {result.paragraphs[0].text}"
    assert result.paragraphs[1].text == "Bu ay icerisinde tum hedefler basariyla tamamlandi ve sistem uretim ortamina deploy edildi.", f"P1 text mismatch"
    assert result.paragraphs[2].text == "Hazirlayan: Omer - 21.05.2026", f"P2 text mismatch"
    print("[OK] Tum metin dogrulamalari basarili!")

if __name__ == "__main__":
    main()
